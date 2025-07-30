from docx import Document
from utils.visual_utils import detect_background_color, detect_checkbox, detect_unit_value

def parse_docx(doc_path):
    doc = Document(doc_path)
    sections = []
    section = {"heading": "", "paragraphs": [], "tables": []}

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.lower().startswith("section"):
            if section["heading"]:
                sections.append(section)
            section = {"heading": text, "paragraphs": [], "tables": []}
        else:
            section["paragraphs"].append(text)

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                value = cell.text.strip()
                color = detect_background_color(cell)
                checkbox = detect_checkbox(cell)
                unit_value = detect_unit_value(cell.text)
                row_data.append({
                    "value": value,
                    "bg_color": color,
                    "checkbox": checkbox,
                    "unit_value": unit_value
                })
            table_data.append(row_data)
        section["tables"].append(table_data)

    sections.append(section)
    return sections
